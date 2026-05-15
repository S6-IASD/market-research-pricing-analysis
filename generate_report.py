from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Styles
    title = doc.add_heading('Rapport Technique et Analytique du Projet Market Research', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Ce rapport présente une analyse approfondie du projet d'analyse de marché et de tarification (Market Research Pricing Analysis). "
        "Il s'agit d'une application full-stack complexe permettant de scraper des données e-commerce, de les analyser via des algorithmes "
        "de Machine Learning (Data Mining) et de les restituer de manière visuelle sur une interface web dynamique."
    )

    doc.add_heading('2. Fonctionnalités Principales', level=1)
    doc.add_paragraph("Le système offre les fonctionnalités suivantes :", style='List Bullet')
    doc.add_paragraph("Recherche et Scraping en direct (Deep Search) de produits sur de grandes plateformes (Jumia, AliExpress, eBay).", style='List Bullet')
    doc.add_paragraph("Analyse avancée de données (Clustering, Détection d'anomalies, Règles d'association) appliquées automatiquement sur les prix extraits.", style='List Bullet')
    doc.add_paragraph("Mise en cache intelligente des résultats d'analyse pour garantir des performances optimales à l'utilisateur.", style='List Bullet')
    doc.add_paragraph("Tableau de bord (Dashboard) interactif générant des visualisations statistiques (BoxPlots de distribution, Histogrammes, Répartitions).", style='List Bullet')

    doc.add_heading('3. Détails Techniques', level=1)

    # Scraping
    doc.add_heading('3.1. Couche Scraping', level=2)
    doc.add_paragraph(
        "La couche de scraping est responsable de la collecte des données brutes en temps réel. Les technologies employées incluent :"
    )
    doc.add_paragraph("Playwright : Indispensable pour interagir avec les sites dynamiques (ex: AliExpress) qui exigent l'exécution de JavaScript pour rendre le contenu.", style='List Bullet')
    doc.add_paragraph("BeautifulSoup4 & lxml : Utilisés pour parser efficacement et rapidement le code HTML statique (ex: Jumia, eBay).", style='List Bullet')
    doc.add_paragraph("Celery (Workers asynchrones) : Étant donné que le scraping est un processus long, il est géré en arrière-plan via des files d'attente pour ne pas bloquer l'interface utilisateur.", style='List Bullet')

    # Data Mining
    doc.add_heading('3.2. Data Mining (Intelligence Artificielle)', level=2)
    doc.add_paragraph(
        "L'intelligence de la plateforme repose sur un pipeline ML. Technologies : Python, Pandas, Scikit-learn, Mlxtend."
    )
    doc.add_paragraph("Nettoyage et Préparation (Pandas) : Transformation des données brutes en DataFrames, filtrage, et génération de variables statistiques (bins d'histogrammes, quartiles).", style='List Bullet')
    doc.add_paragraph("Clustering (K-Means) : Segmentation non supervisée des produits en 3 gammes (Entrée, Milieu, Haut de gamme) basée sur les vecteurs de prix.", style='List Bullet')
    doc.add_paragraph("Détection d'Anomalies (Isolation Forest + LOF) : Combinaison de deux algorithmes de détection (par isolation et par densité locale) pour identifier de fausses annonces ou des prix aberrants.", style='List Bullet')
    doc.add_paragraph("Règles d'Association (FP-Growth) : Extraction de règles métiers et d'insights (ex: Si l'utilisateur est sur Jumia, alors le produit a 80% de chance d'être en Entrée de gamme).", style='List Bullet')

    # Backend
    doc.add_heading('3.3. Backend', level=2)
    doc.add_paragraph(
        "Le cœur applicatif est une API REST robuste et scalable. Technologies : Django, Django REST Framework, PostgreSQL, Redis."
    )
    doc.add_paragraph("Architecture MVT/REST : Modélisation forte des entités (Product, PriceSnapshot, AnalysisResult, SearchTask).", style='List Bullet')
    doc.add_paragraph("Base de données (PostgreSQL) : Utilisation du composant 'SearchVector' pour permettre une recherche 'Full-Text' rapide et optimisée sur les mots-clés et catégories.", style='List Bullet')
    doc.add_paragraph("Redis (In-Memory Datastore) : Gère le cache des limites de requêtes (Rate Limiting 5 requêtes/jour/utilisateur) et sert de Message Broker principal pour Celery.", style='List Bullet')
    doc.add_paragraph("Système de Cache Analytique : Le modèle AnalysisResult stocke les objets JSON lourds des résultats Data Mining pendant 24h, offrant des réponses sub-millisecondes pour les requêtes répétées.", style='List Bullet')

    # Frontend
    doc.add_heading('3.4. Frontend', level=2)
    doc.add_paragraph(
        "L'interface client est une Single Page Application (SPA). Technologies : React, Vite, TypeScript, Ant Design, Recharts."
    )
    doc.add_paragraph("Framework UI (Ant Design) : Utilisation de composants d'entreprise robustes (Cards, List, Badges, Layout) assurant un design professionnel et responsive.", style='List Bullet')
    doc.add_paragraph("Visualisation des Données (Recharts) : Implémentation graphique (ComposedChart pour simuler les BoxPlots par quartile, BarChart, PieChart) s'alimentant en temps réel de l'API.", style='List Bullet')
    doc.add_paragraph("Architecture : Utilisation intensive des Hooks React (useEffect, useState), avec gestion d'états asynchrones (Loaders) pour refléter l'évolution des tâches de fond.", style='List Bullet')

    doc.add_heading('4. Problèmes et Défis Rencontrés', level=1)
    
    p1 = doc.add_paragraph("1. Logique d'Extraction de Mots-Clés : ", style='List Bullet')
    p1.add_run("Un problème identifié résidait dans l'extracteur de mots-clés (`keyword_extractor.py`) qui supprimait automatiquement les mots de 2 lettres ou moins. Ainsi, une recherche précise comme 'laptop 16 gb ram' était amputée et devenait 'laptop ram', impactant la précision du scraping et biaisant l'analyse finale. Cela nécessite une adaptation pour conserver les valeurs numériques et unités techniques.")
    
    p2 = doc.add_paragraph("2. Lourdeur de l'Infrastructure Docker : ", style='List Bullet')
    p2.add_run("Le conteneur backend embarque l'ensemble du stack ML (Scikit-learn, Pandas) ainsi que Playwright et ses dépendances navigateurs. La taille massive de l'image (plusieurs Go) provoque des saturations de stockage et des arrêts inopinés du daemon Docker (Erreur EOF) lors de l'étape de construction (build) sur des environnements limités en ressources.")
    
    p3 = doc.add_paragraph("3. Instabilité du Scraping Live : ", style='List Bullet')
    p3.add_run("Le scraping en temps réel est fondamentalement lent et soumis aux contremesures anti-bots des grandes plateformes. Ce défi a contraint l'architecture à évoluer vers un système 100% asynchrone (Celery) couplé à une base de données de sauvegarde, garantissant que l'utilisateur n'ait pas à subir le temps d'exécution (timeout).")

    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "Le projet 'Market Research Pricing Analysis' est une plateforme d'intelligence de données complète. "
        "L'intégration réussie d'un pipeline de Data Mining hors-ligne vers une architecture backend Django dynamique, "
        "couplée à une interface React réactive, positionne ce projet comme un outil analytique de pointe. "
        "La gestion des défis d'infrastructure asynchrone démontre une excellente maîtrise des enjeux d'ingénierie logicielle."
    )

    doc.save('Rapport_Technique_Projet.docx')
    print("Document Word généré avec succès : Rapport_Technique_Projet.docx")

if __name__ == '__main__':
    main()
